from docx import Document
from docx.shared import Inches


document = Document()
# title
head = input('Enter the title: ')
subtitle = input('Subtitle(your position): ')
document.add_heading(head, 0)
document.add_heading(f'{subtitle}')

# profile picture
document.add_picture(
    'me.png',
    width=Inches(1.5)
)

# name, phone, email details
first_name = input('Enter your name: ')
last_name = input('Enter your last name: ')
phone = input('Enter your phone number: ')
email = input('Enter your email: ')
document.add_heading('INFO')
document.add_paragraph(f'{first_name} {last_name}\nPhone: {phone}\nEmail: {email}').bold = True


# education
document.add_heading('EDUCATION')
p = document.add_paragraph()

education = input('Enter the name of your institution: ')
from_date = input('From date: ')
to_date = input('To date: ')

p.add_run(education + ' ').bold = True
p.add_run(from_date + '-' + to_date).italic = True


# more education
while True:
    has_more_education = input('Do you have any other education? Yes or No ')
    if has_more_education.lower() == 'yes':
        p = document.add_paragraph()

        education = input('Enter the name of your institution: ')
        from_date = input('From date: ')
        to_date = input('To date: ')

        p.add_run(education + ' ').bold = True
        p.add_run(from_date + '-' + to_date).italic = True
    else:
        break


# work experience
document.add_heading('WORK EXPERIENCE')
p = document.add_paragraph()

company = input('Enter company name: ')
from_date = input('From date: ')
to_date = input('To date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    f'Describe your experience in {company}: '
)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experience = input('Do you have more experience? Yes or No ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company name: ')
        from_date = input('From date: ')
        to_date = input('To date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            f'Describe your experience in {company}: '
        )
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('SKILLS')
skill = input('Enter skill ')
p = document.add_paragraph(skill)

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
    else:
        break

# languages
document.add_heading('LANGUAGES')
language = input('Enter the languages you know(and your level): ')
p = document.add_paragraph(language)
p.style = 'List Bullet'

while True:
    has_more_language = input('Do you know any other languages? Yes or No ')
    if has_more_language.lower() == 'yes':
        language = input('Enter language ')
        p = document.add_paragraph(language)
        p.style = 'List Bullet'
    else:
        break

# about me
document.add_heading('PROFILE')
document.add_paragraph(
    input('Tell me about yourself: ')
)


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Vladyslav Baisha code."

document.save('cv.docx')
